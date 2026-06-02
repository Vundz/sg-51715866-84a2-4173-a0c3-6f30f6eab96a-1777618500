"""
Generate Nursery App Process Flow Document
Produces Process_Flows.docx with matplotlib flowcharts embedded as images.
"""

import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colour palette ──────────────────────────────────────────────────────────
C = {
    "start":    "#2E7D32",   # dark green
    "end":      "#B71C1C",   # dark red
    "process":  "#1565C0",   # dark blue
    "decision": "#E65100",   # dark orange
    "data":     "#4A148C",   # dark purple
    "system":   "#00695C",   # teal
    "white":    "#FFFFFF",
    "bg":       "#F8F9FA",
}

# ─── Low-level drawing helpers ────────────────────────────────────────────────

def _box(ax, x, y, w, h, color, text, fontsize=7.5, text_color="white",
         shape="rect", zorder=3):
    """Draw a labelled box; shape: rect | diamond | rounded | parallelogram"""
    if shape == "diamond":
        dx, dy = w / 2, h / 2
        poly = plt.Polygon(
            [[x, y + dy], [x + dx, y], [x + w, y + dy], [x + dx, y + h]],
            closed=True, facecolor=color, edgecolor="white", linewidth=1.2,
            zorder=zorder)
        ax.add_patch(poly)
        cx, cy = x + dx, y + dy
    elif shape == "rounded":
        patch = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.05",
            facecolor=color, edgecolor="white", linewidth=1.2, zorder=zorder)
        ax.add_patch(patch)
        cx, cy = x + w / 2, y + h / 2
    elif shape == "parallelogram":
        skew = h * 0.3
        poly = plt.Polygon(
            [[x + skew, y], [x + w + skew, y],
             [x + w, y + h], [x, y + h]],
            closed=True, facecolor=color, edgecolor="white", linewidth=1.2,
            zorder=zorder)
        ax.add_patch(poly)
        cx, cy = x + w / 2 + skew / 2, y + h / 2
    else:
        patch = FancyBboxPatch(
            (x, y), w, h, boxstyle="square,pad=0",
            facecolor=color, edgecolor="white", linewidth=1.2, zorder=zorder)
        ax.add_patch(patch)
        cx, cy = x + w / 2, y + h / 2

    ax.text(cx, cy, text, ha="center", va="center",
            fontsize=fontsize, color=text_color, fontweight="bold",
            wrap=True, zorder=zorder + 1,
            multialignment="center")
    return cx, cy


def _arrow(ax, x1, y1, x2, y2, label="", color="#555555"):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color,
                                lw=1.4, mutation_scale=12),
                zorder=2)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx + 0.05, my, label, fontsize=6.5, color=color,
                va="center", style="italic")


def _save(fig, ax):
    ax.set_aspect("equal")
    ax.axis("off")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight",
                facecolor=C["bg"])
    buf.seek(0)
    plt.close(fig)
    return buf


# ─── Individual flow diagrams ─────────────────────────────────────────────────

def flow_onboarding():
    """System Setup & Onboarding"""
    fig, ax = plt.subplots(figsize=(9, 11))
    ax.set_xlim(0, 10); ax.set_ylim(0, 12)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 1 – System Setup & Onboarding",
                 fontsize=11, fontweight="bold", pad=8)

    steps = [
        (4.0, 10.8, 2.0, 0.6, C["start"],   "START",                   "rounded"),
        (4.0,  9.8, 2.0, 0.7, C["process"], "Admin Login",             "rect"),
        (4.0,  8.8, 2.0, 0.7, C["process"], "Create Locations\n(greenhouses/areas)", "rect"),
        (4.0,  7.8, 2.0, 0.7, C["process"], "Create Plant Types\n& Varieties",       "rect"),
        (4.0,  6.8, 2.0, 0.7, C["process"], "Set Up Inventory\nItems & Categories",  "rect"),
        (4.0,  5.8, 2.0, 0.7, C["process"], "Configure Scouting\nPest/Disease Types","rect"),
        (4.0,  4.8, 2.0, 0.7, C["process"], "Create BOM Templates\n(Production)",    "rect"),
        (3.5,  3.6, 3.0, 0.9, C["decision"],"All reference data\ncomplete?",         "diamond"),
        (4.0,  2.5, 2.0, 0.7, C["process"], "Create User Accounts\n& Assign Roles",  "rect"),
        (4.0,  1.5, 2.0, 0.7, C["process"], "Set Role Permissions\nper Module",       "rect"),
        (4.0,  0.3, 2.0, 0.7, C["start"],   "SYSTEM READY",            "rounded"),
    ]

    centers = []
    for x, y, w, h, col, txt, shp in steps:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        centers.append((cx, cy, y, h))

    # straight arrows
    for i in range(len(centers) - 1):
        if i == 7:   # skip; handled separately
            continue
        _arrow(ax, centers[i][0], centers[i][2],
                   centers[i+1][0], centers[i+1][2] + centers[i+1][3])

    # decision → yes (down)
    dec_cx, dec_cy = centers[7][0], centers[7][2] + centers[7][3]/2
    _arrow(ax, dec_cx, centers[7][2],
               centers[8][0], centers[8][2] + centers[8][3], label="Yes")

    # decision → No (loop back to step 2)
    _arrow(ax, dec_cx + 1.5, dec_cy,
               7.8, 10.15, color="#E65100")
    ax.annotate("", xy=(6.0, 10.15), xytext=(7.8, 10.15),
                arrowprops=dict(arrowstyle="-|>", color="#E65100", lw=1.4))
    ax.text(8.0, (dec_cy + 10.15) / 2, "No – fill\nmissing data",
            fontsize=6.5, color="#E65100", style="italic")

    return _save(fig, ax)


def flow_planting():
    """Planting Lifecycle"""
    fig, ax = plt.subplots(figsize=(11, 13))
    ax.set_xlim(0, 12); ax.set_ylim(0, 14)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 2 – Planting Lifecycle",
                 fontsize=11, fontweight="bold", pad=8)

    # Main spine x=5
    spine = [
        (4.5, 13.0, 3, 0.6,  C["start"],    "START",                       "rounded"),
        (4.5, 11.9, 3, 0.8,  C["process"],  "Select Plant Type\n& Location","rect"),
        (4.0, 10.7, 4, 0.9,  C["decision"], "Seed in\nInventory?",          "diamond"),
        (4.5,  9.5, 3, 0.8,  C["process"],  "Deduct Seed Quantity\nfrom Inventory","rect"),
        (4.5,  8.4, 3, 0.8,  C["process"],  "Auto-generate\nBatch Number",  "rect"),
        (4.5,  7.3, 3, 0.8,  C["process"],  "Set Planting Date\n& Expected Harvest","rect"),
        (4.5,  6.2, 3, 0.8,  C["process"],  "Save Planting Record",         "rect"),
        (4.0,  5.0, 4, 0.9,  C["decision"], "Reservation exists\nfor this batch?","diamond"),
        (4.5,  3.9, 3, 0.8,  C["system"],   "Link Reservation\nto Planting","rect"),
        (4.5,  2.8, 3, 0.8,  C["process"],  "Monitor Growth\n(Scouting / Treatments)","rect"),
        (4.5,  1.7, 3, 0.8,  C["process"],  "Planting Status:\nActive → Harvested","rect"),
        (4.5,  0.4, 3, 0.6,  C["end"],      "END",                          "rounded"),
    ]

    cx_list = []
    for x, y, w, h, col, txt, shp in spine:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        cx_list.append((cx, y, h))

    for i in range(len(cx_list) - 1):
        if i in (2, 7):
            continue
        _arrow(ax, cx_list[i][0], cx_list[i][1],
                   cx_list[i+1][0], cx_list[i+1][1] + cx_list[i+1][2])

    # decision 1 (seed in inventory) → yes → deduct
    dec1_cx = cx_list[2][0]
    _arrow(ax, dec1_cx, cx_list[2][1], cx_list[3][0],
               cx_list[3][1] + cx_list[3][2], label="Yes")
    # decision 1 → no → skip deduct (bypass to batch number)
    _arrow(ax, dec1_cx + 2.0, cx_list[2][1] + spine[2][3]/2,
               9.5, spine[4][1] + spine[4][3]/2, color="#888")
    ax.annotate("", xy=(7.5, spine[4][1] + spine[4][3]/2),
                xytext=(9.5, spine[4][1] + spine[4][3]/2),
                arrowprops=dict(arrowstyle="-|>", color="#888", lw=1.2))
    ax.text(10.5, (cx_list[2][1] + spine[4][1])/2, "No – no\nseed tracking",
            fontsize=6.5, color="#888", style="italic", va="center")

    # decision 2 (reservation exists) → yes → link
    dec2_cx = cx_list[7][0]
    _arrow(ax, dec2_cx, cx_list[7][1], cx_list[8][0],
               cx_list[8][1] + cx_list[8][2], label="Yes")
    # decision 2 → no → skip link
    _arrow(ax, dec2_cx + 2.0, cx_list[7][1] + spine[7][3]/2,
               9.5, spine[9][1] + spine[9][3]/2, color="#888")
    ax.annotate("", xy=(7.5, spine[9][1] + spine[9][3]/2),
                xytext=(9.5, spine[9][1] + spine[9][3]/2),
                arrowprops=dict(arrowstyle="-|>", color="#888", lw=1.2))
    ax.text(10.5, (cx_list[7][1] + spine[9][1])/2, "No",
            fontsize=6.5, color="#888", style="italic", va="center")

    # reservation link → monitor
    _arrow(ax, cx_list[8][0], cx_list[8][1], cx_list[9][0],
               cx_list[9][1] + cx_list[9][2])

    return _save(fig, ax)


def flow_harvest_dispatch():
    """Harvest & Dispatch Slip Flow"""
    fig, ax = plt.subplots(figsize=(13, 14))
    ax.set_xlim(0, 14); ax.set_ylim(0, 15)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 3 – Harvest & Dispatch Slip",
                 fontsize=11, fontweight="bold", pad=8)

    # ── Harvest path (left lane x≈1) ─────────────────────────
    ax.text(2.5, 14.7, "HARVEST PATH", ha="center", fontsize=8,
            fontweight="bold", color=C["process"])
    h = [
        (1, 13.8, 3, 0.65, C["start"],    "Create Harvest",         "rounded"),
        (1, 12.8, 3, 0.8,  C["process"],  "Select Planting Batch",  "rect"),
        (0.5, 11.6, 4, 0.9, C["decision"],"Qty ≤ Remaining\nQuantity?","diamond"),
        (1, 10.5, 3, 0.8,  C["process"],  "Check Reservation\nConflicts","rect"),
        (0.5, 9.3, 4, 0.9,  C["decision"],"Reservation\nConflict?",  "diamond"),
        (1,  8.2, 3, 0.8,  C["system"],   "Warn User – Override\nor Cancel?","rect"),
        (1,  7.1, 3, 0.8,  C["process"],  "Record Harvest\n(qty, quality, status)","rect"),
        (1,  6.0, 3, 0.8,  C["system"],   "Update Planting\nremaining_quantity","rect"),
        (1,  4.9, 3, 0.8,  C["process"],  "Print Harvest\nRecord (optional)","rect"),
        (1,  3.7, 3, 0.6,  C["end"],      "HARVEST SAVED",          "rounded"),
    ]
    hc = []
    for x, y, w, h2, col, txt, shp in h:
        cx, cy = _box(ax, x, y, w, h2, col, txt, shape=shp)
        hc.append((cx, y, h2))

    for i in range(len(hc)-1):
        if i in (2, 4):
            continue
        _arrow(ax, hc[i][0], hc[i][1], hc[i+1][0], hc[i+1][1]+hc[i+1][2])

    # qty decision yes/no
    _arrow(ax, hc[2][0], hc[2][1], hc[3][0], hc[3][1]+hc[3][2], label="Yes")
    ax.text(-0.3, hc[2][1]+0.45, "No →\nInsufficient",
            fontsize=6, color="red", style="italic")

    # reservation conflict yes → warn
    _arrow(ax, hc[4][0], hc[4][1], hc[5][0], hc[5][1]+hc[5][2], label="Yes")
    # no conflict → direct to record
    _arrow(ax, hc[4][0]+2.0, hc[4][1]+h[4][3]/2,
               4.8, hc[6][1]+hc[6][2], color="#888")
    ax.annotate("", xy=(4.0, hc[6][1]+hc[6][2]),
                xytext=(4.8, hc[6][1]+hc[6][2]),
                arrowprops=dict(arrowstyle="-|>", color="#888", lw=1.2))
    ax.text(5.0, (hc[4][1]+hc[6][1])/2, "No", fontsize=6.5,
            color="#888", style="italic", va="center")

    # warn → record (override)
    _arrow(ax, hc[5][0], hc[5][1], hc[6][0], hc[6][1]+hc[6][2], label="Override")

    # ── Dispatch Slip path (right lane x≈8) ──────────────────
    ax.text(10.5, 14.7, "DISPATCH SLIP PATH", ha="center", fontsize=8,
            fontweight="bold", color=C["data"])
    d = [
        (8, 13.8, 3, 0.65, C["start"],    "Create Dispatch Slip",   "rounded"),
        (8, 12.8, 3, 0.8,  C["process"],  "Select Planting Batch\n& Enter Qty","rect"),
        (7.5, 11.6, 4, 0.9, C["decision"],"Stock Available\n(remaining ≥ qty)?","diamond"),
        (8, 10.5, 3, 0.8,  C["system"],   "Auto-Create Harvest\nRecord",       "rect"),
        (8,  9.4, 3, 0.8,  C["system"],   "Update remaining_\nquantity",       "rect"),
        (8,  8.3, 3, 0.8,  C["process"],  "Save Slip as\nFULFILLED",           "rect"),
        (8,  7.1, 3, 0.8,  C["process"],  "Save Slip as\nPENDING",             "rect"),
        (7.5, 5.9, 4, 0.9, C["decision"], "Stock replenished\nlater?",         "diamond"),
        (8,  4.8, 3, 0.8,  C["process"],  "Edit Slip &\nFulfill Now",           "rect"),
        (8,  3.7, 3, 0.8,  C["process"],  "Print Dispatch\nSlip (80mm)",        "rect"),
        (8,  2.6, 3, 0.6,  C["end"],      "DISPATCH SLIP DONE",     "rounded"),
    ]
    dc = []
    for x, y, w, h2, col, txt, shp in d:
        cx, cy = _box(ax, x, y, w, h2, col, txt, shape=shp)
        dc.append((cx, y, h2))

    for i in range(len(dc)-1):
        if i in (2, 7):
            continue
        _arrow(ax, dc[i][0], dc[i][1], dc[i+1][0], dc[i+1][1]+dc[i+1][2])

    # stock available yes → auto harvest
    _arrow(ax, dc[2][0], dc[2][1], dc[3][0], dc[3][1]+dc[3][2], label="Yes")
    # stock not available → pending
    _arrow(ax, dc[2][0]-2.0, dc[2][1]+d[2][3]/2,
               6.5, dc[6][1]+dc[6][2], color="#E65100")
    ax.annotate("", xy=(8.0, dc[6][1]+dc[6][2]),
                xytext=(6.5, dc[6][1]+dc[6][2]),
                arrowprops=dict(arrowstyle="-|>", color="#E65100", lw=1.4))
    ax.text(5.0, (dc[2][1]+dc[6][1])/2, "No",
            fontsize=6.5, color="#E65100", style="italic", va="center")

    # fulfilled → print (skip pending path)
    _arrow(ax, dc[5][0], dc[5][1], dc[9][0], dc[9][1]+dc[9][2],
           color="#1565C0")
    ax.text(11.5, (dc[5][1]+dc[9][1])/2, "→ Print",
            fontsize=6, color="#1565C0", style="italic", va="center")

    # pending → replenished yes → edit & fulfill
    _arrow(ax, dc[7][0], dc[7][1], dc[8][0], dc[8][1]+dc[8][2], label="Yes")
    # pending → no replenishment → cancel option
    ax.text(12.5, dc[7][1]+0.45, "No →\nCancel Slip",
            fontsize=6, color="red", style="italic", ha="center")

    # edit & fulfill → print
    _arrow(ax, dc[8][0], dc[8][1], dc[9][0], dc[9][1]+dc[9][2])

    # lane divider
    ax.axvline(7.3, color="#CCCCCC", lw=1, linestyle="--", zorder=1)

    return _save(fig, ax)


def flow_reservation():
    """Reservation Flow"""
    fig, ax = plt.subplots(figsize=(10, 13))
    ax.set_xlim(0, 11); ax.set_ylim(0, 14)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 4 – Reservation Management",
                 fontsize=11, fontweight="bold", pad=8)

    steps = [
        (4.0, 13.1, 3, 0.6,  C["start"],    "START",                       "rounded"),
        (4.0, 12.0, 3, 0.8,  C["process"],  "Customer Requests\nSeedlings", "rect"),
        (4.0, 10.9, 3, 0.8,  C["process"],  "Select Planting Batch(es)",    "rect"),
        (3.5,  9.7, 4, 0.9,  C["decision"], "Batch has enough\navailable qty?","diamond"),
        (4.0,  8.6, 3, 0.8,  C["process"],  "Enter Qty, Date,\nCustomer Info","rect"),
        (3.5,  7.4, 4, 0.9,  C["decision"], "Multiple Batches?",             "diamond"),
        (4.0,  6.3, 3, 0.8,  C["system"],   "Create Reservation\nper Batch", "rect"),
        (4.0,  5.2, 3, 0.8,  C["process"],  "Set Payment Status\n(pending/partial/paid)","rect"),
        (3.5,  4.0, 4, 0.9,  C["decision"], "Ready to\nFulfil?",             "diamond"),
        (4.0,  2.9, 3, 0.8,  C["process"],  "Mark Completed –\nEnter Actual Qty","rect"),
        (4.0,  1.8, 3, 0.8,  C["system"],   "Deduct from Planting\nremaining_quantity","rect"),
        (4.0,  0.5, 3, 0.6,  C["end"],      "RESERVATION CLOSED",           "rounded"),
    ]

    cx_list = []
    for x, y, w, h, col, txt, shp in steps:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        cx_list.append((cx, y, h))

    for i in range(len(cx_list)-1):
        if i in (3, 5, 8):
            continue
        _arrow(ax, cx_list[i][0], cx_list[i][1], cx_list[i+1][0],
               cx_list[i+1][1]+cx_list[i+1][2])

    # batch qty decision
    _arrow(ax, cx_list[3][0], cx_list[3][1], cx_list[4][0],
               cx_list[4][1]+cx_list[4][2], label="Yes")
    ax.text(0.2, cx_list[3][1]+0.45, "No → choose\nanother batch",
            fontsize=6, color="red", style="italic")

    # multiple batches → yes
    _arrow(ax, cx_list[5][0], cx_list[5][1], cx_list[6][0],
               cx_list[6][1]+cx_list[6][2], label="Yes")
    # multiple batches → no (single reservation)
    _arrow(ax, cx_list[5][0]+2.0, cx_list[5][1]+steps[5][3]/2,
               9.5, cx_list[6][1]+cx_list[6][2], color="#888")
    ax.annotate("", xy=(7.0, cx_list[6][1]+cx_list[6][2]),
                xytext=(9.5, cx_list[6][1]+cx_list[6][2]),
                arrowprops=dict(arrowstyle="-|>", color="#888", lw=1.2))
    ax.text(9.8, (cx_list[5][1]+cx_list[6][1])/2, "No",
            fontsize=6.5, color="#888", style="italic", va="center")

    # ready to fulfil
    _arrow(ax, cx_list[8][0], cx_list[8][1], cx_list[9][0],
               cx_list[9][1]+cx_list[9][2], label="Yes")
    # cancel path
    ax.text(0.2, cx_list[8][1]+0.45, "No → Cancel\n(return stock)",
            fontsize=6, color="red", style="italic")

    return _save(fig, ax)


def flow_inventory():
    """Inventory Management Flow"""
    fig, ax = plt.subplots(figsize=(11, 12))
    ax.set_xlim(0, 12); ax.set_ylim(0, 13)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 5 – Inventory Management",
                 fontsize=11, fontweight="bold", pad=8)

    # Two sub-flows: Manual and Auto-triggered
    ax.text(2.5, 12.7, "MANUAL TRANSACTIONS", ha="center", fontsize=8,
            fontweight="bold", color=C["process"])
    ax.text(8.5, 12.7, "AUTO-TRIGGERED", ha="center", fontsize=8,
            fontweight="bold", color=C["system"])

    # Manual path
    m = [
        (1, 12.0, 3, 0.6,  C["start"],    "Staff / Manager",        "rounded"),
        (1, 11.0, 3, 0.8,  C["process"],  "Select Inventory Item",  "rect"),
        (0.5, 9.8, 4, 0.9, C["decision"], "Transaction\nType?",     "diamond"),
        (1,  8.7, 3, 0.8,  C["process"],  "Purchase:\nAdd qty + unit price","rect"),
        (1,  7.6, 3, 0.8,  C["process"],  "Usage:\nDeduct qty used","rect"),
        (1,  6.5, 3, 0.8,  C["process"],  "Adjustment:\nAdd or subtract","rect"),
        (1,  5.4, 3, 0.8,  C["process"],  "Waste:\nDeduct waste qty","rect"),
        (1,  4.3, 3, 0.8,  C["system"],   "Update current_stock\n& log transaction","rect"),
        (0.5, 3.1,4, 0.9,  C["decision"], "Stock below\nminimum?",  "diamond"),
        (1,  2.0, 3, 0.8,  C["process"],  "Show Low Stock\nAlert",  "rect"),
        (1,  0.8, 3, 0.6,  C["end"],      "TRANSACTION SAVED",      "rounded"),
    ]
    mc = []
    for x, y, w, h, col, txt, shp in m:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        mc.append((cx, y, h))

    # straight arrows (skip decisions)
    for i in [0,1,7,8,9]:
        if i+1 < len(mc):
            _arrow(ax, mc[i][0], mc[i][1], mc[i+1][0], mc[i+1][1]+mc[i+1][2])

    # decision branches
    dec_cx = mc[2][0]
    for idx, label in [(3,"Purchase"), (4,"Usage"), (5,"Adjust"), (6,"Waste")]:
        _arrow(ax, dec_cx, mc[2][1], mc[idx][0], mc[idx][1]+mc[idx][2], label=label)
    # all 4 converge to update
    for idx in [3,4,5,6]:
        _arrow(ax, mc[idx][0], mc[idx][1], mc[7][0], mc[7][1]+mc[7][2])

    # low stock yes → alert
    _arrow(ax, mc[8][0], mc[8][1], mc[9][0], mc[9][1]+mc[9][2], label="Yes")
    ax.text(0.1, mc[8][1]+0.45, "No", fontsize=6.5, color="#888",
            style="italic")
    _arrow(ax, mc[9][0], mc[9][1], mc[10][0], mc[10][1]+mc[10][2])

    # Auto path (right side)
    a = [
        (8, 12.0, 3, 0.6,  C["start"],    "Planting Created",        "rounded"),
        (8, 11.0, 3, 0.8,  C["system"],   "System checks for\nSeed inventory item","rect"),
        (7.5, 9.8, 4, 0.9, C["decision"], "Seed item\nexists?",      "diamond"),
        (8,  8.7, 3, 0.8,  C["system"],   "Auto-create Usage\nTransaction",       "rect"),
        (8,  7.6, 3, 0.8,  C["system"],   "Deduct seed qty\nfrom current_stock",  "rect"),
        (8,  6.5, 3, 0.6,  C["end"],      "STOCK UPDATED",           "rounded"),
    ]
    ac = []
    for x, y, w, h, col, txt, shp in a:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        ac.append((cx, y, h))

    for i in range(len(ac)-1):
        if i == 2:
            continue
        _arrow(ax, ac[i][0], ac[i][1], ac[i+1][0], ac[i+1][1]+ac[i+1][2])
    _arrow(ax, ac[2][0], ac[2][1], ac[3][0], ac[3][1]+ac[3][2], label="Yes")
    ax.text(12.2, ac[2][1]+0.45, "No", fontsize=6.5, color="#888",
            style="italic")

    ax.axvline(6.5, color="#CCCCCC", lw=1, linestyle="--", zorder=1)

    return _save(fig, ax)


def flow_scouting_treatment():
    """Scouting → Treatment Flow"""
    fig, ax = plt.subplots(figsize=(11, 13))
    ax.set_xlim(0, 12); ax.set_ylim(0, 14)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 6 – Scouting & Treatment",
                 fontsize=11, fontweight="bold", pad=8)

    steps = [
        (4.5, 13.1, 3, 0.6,  C["start"],    "Scout visits nursery",         "rounded"),
        (4.5, 12.0, 3, 0.8,  C["process"],  "Create Scouting Report\n(date, observer, batch)","rect"),
        (4.5, 10.9, 3, 0.8,  C["process"],  "Log Observations:\nPests / Diseases / Nutrients","rect"),
        (4.5,  9.8, 3, 0.8,  C["process"],  "Set Severity Levels\n& Record Evidence","rect"),
        (4.0,  8.6, 4, 0.9,  C["decision"], "Issues identified\nrequire treatment?","diamond"),
        (4.5,  7.5, 3, 0.8,  C["process"],  "Create Treatment Record\n(chemical, method, date)","rect"),
        (4.0,  6.3, 4, 0.9,  C["decision"], "Apply to single\nor bulk batches?","diamond"),
        (1.5,  5.2, 3, 0.8,  C["process"],  "Single Batch:\nLink to planting","rect"),
        (7.0,  5.2, 3, 0.8,  C["process"],  "Bulk Apply:\nSelect by location","rect"),
        (4.5,  4.1, 3, 0.8,  C["system"],   "Create planting_\ntreatments records","rect"),
        (4.5,  3.0, 3, 0.8,  C["process"],  "Record who applied\n& quantity used","rect"),
        (4.5,  1.9, 3, 0.8,  C["process"],  "Deduct chemical\nfrom Inventory","rect"),
        (4.5,  0.7, 3, 0.6,  C["end"],      "TREATMENT COMPLETE",           "rounded"),
    ]

    cx_list = []
    for x, y, w, h, col, txt, shp in steps:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        cx_list.append((cx, y, h))

    skip = {4, 6}
    for i in range(len(cx_list)-1):
        if i in skip or i in {7, 8}:
            continue
        _arrow(ax, cx_list[i][0], cx_list[i][1], cx_list[i+1][0],
               cx_list[i+1][1]+cx_list[i+1][2])

    # issues decision → yes
    _arrow(ax, cx_list[4][0], cx_list[4][1], cx_list[5][0],
               cx_list[5][1]+cx_list[5][2], label="Yes")
    ax.text(9.0, cx_list[4][1]+0.45, "No → Report\nfiled only",
            fontsize=6, color="#888", style="italic")

    # bulk vs single
    _arrow(ax, cx_list[6][0]-2.0, cx_list[6][1]+steps[6][3]/2,
               cx_list[7][0], cx_list[7][1]+cx_list[7][2], label="Single")
    _arrow(ax, cx_list[6][0]+2.0, cx_list[6][1]+steps[6][3]/2,
               cx_list[8][0], cx_list[8][1]+cx_list[8][2], label="Bulk")

    # both converge to planting_treatments
    _arrow(ax, cx_list[7][0], cx_list[7][1], cx_list[9][0],
               cx_list[9][1]+cx_list[9][2])
    _arrow(ax, cx_list[8][0], cx_list[8][1], cx_list[9][0],
               cx_list[9][1]+cx_list[9][2])

    return _save(fig, ax)


def flow_bom_production():
    """BOM / Production Cost Flow"""
    fig, ax = plt.subplots(figsize=(10, 13))
    ax.set_xlim(0, 11); ax.set_ylim(0, 14)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 7 – BOM & Production Cost Tracking",
                 fontsize=11, fontweight="bold", pad=8)

    steps = [
        (4.0, 13.1, 3, 0.6,  C["start"],    "Manager/Admin",                "rounded"),
        (4.0, 12.0, 3, 0.8,  C["process"],  "Create BOM Template\n(product type)","rect"),
        (4.0, 10.9, 3, 0.8,  C["process"],  "Add BOM Items\n(inventory items or custom)","rect"),
        (4.0,  9.8, 3, 0.8,  C["process"],  "Set Qty Formulas\n(per seedling)","rect"),
        (4.0,  8.7, 3, 0.8,  C["process"],  "Assign to Planting\nBatch",     "rect"),
        (3.5,  7.5, 4, 0.9,  C["decision"], "Calculate\nproduction cost?",   "diamond"),
        (4.0,  6.4, 3, 0.8,  C["system"],   "Pull current prices\nfrom Inventory","rect"),
        (4.0,  5.3, 3, 0.8,  C["system"],   "Calculate total cost\nper batch","rect"),
        (4.0,  4.2, 3, 0.8,  C["process"],  "Store in\nplanting_bom_costs",  "rect"),
        (4.0,  3.1, 3, 0.8,  C["process"],  "View Profit Analysis\n(cost vs selling price)","rect"),
        (4.0,  2.0, 3, 0.8,  C["process"],  "Use Chemical Calculator\nfor mixing ratios","rect"),
        (4.0,  0.8, 3, 0.6,  C["end"],      "PRODUCTION COSTED",            "rounded"),
    ]

    cx_list = []
    for x, y, w, h, col, txt, shp in steps:
        cx, cy = _box(ax, x, y, w, h, col, txt, shape=shp)
        cx_list.append((cx, y, h))

    for i in range(len(cx_list)-1):
        if i == 5:
            continue
        _arrow(ax, cx_list[i][0], cx_list[i][1], cx_list[i+1][0],
               cx_list[i+1][1]+cx_list[i+1][2])

    _arrow(ax, cx_list[5][0], cx_list[5][1], cx_list[6][0],
               cx_list[6][1]+cx_list[6][2], label="Yes")
    ax.text(9.0, cx_list[5][1]+0.45, "No → Skip\ncost calc",
            fontsize=6, color="#888", style="italic")

    return _save(fig, ax)


def flow_permissions():
    """Role-Based Access Flow"""
    fig, ax = plt.subplots(figsize=(12, 9))
    ax.set_xlim(0, 13); ax.set_ylim(0, 10)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 8 – Role-Based Access Control",
                 fontsize=11, fontweight="bold", pad=8)

    # Roles as swim lanes
    roles = ["Admin", "Manager", "Staff", "Viewer"]
    role_colors = ["#1565C0", "#2E7D32", "#E65100", "#4A148C"]
    lane_y = [8.5, 6.5, 4.5, 2.5]

    for i, (role, color, y) in enumerate(zip(roles, role_colors, lane_y)):
        ax.axhline(y+1.8, color="#DDDDDD", lw=0.8, zorder=0)
        ax.text(0.3, y+0.9, role, fontsize=9, fontweight="bold",
                color=color, va="center",
                bbox=dict(boxstyle="round,pad=0.3", facecolor=color,
                          edgecolor="white", alpha=0.15))

    modules = ["Locations", "Plant Types", "Plantings", "Harvests",
               "Treatments", "Reservations", "Inventory", "Reports", "Admin"]
    perms = {
        "Admin":   ["CRUD"]*8 + ["CRUD"],
        "Manager": ["CRU","CRU","CRU","CRU","CRU","CRU","CRU","R","—"],
        "Staff":   ["R","R","CRU","CRU","CRU","CRU","CRU","R","—"],
        "Viewer":  ["R","R","R","R","R","R","R","R","—"],
    }
    perm_colors = {
        "CRUD": "#1B5E20", "CRU": "#1565C0", "R": "#E65100", "—": "#9E9E9E"
    }

    col_x = [1.8 + i * 1.2 for i in range(len(modules))]
    for j, mod in enumerate(modules):
        ax.text(col_x[j], 9.7, mod, fontsize=6.5, ha="center", va="center",
                rotation=30, fontweight="bold", color="#333333")

    for i, (role, y) in enumerate(zip(roles, lane_y)):
        for j, mod in enumerate(modules):
            p = perms[role][j]
            color = perm_colors[p]
            _box(ax, col_x[j]-0.45, y+0.2, 0.9, 0.8,
                 color, p, fontsize=7, shape="rounded")

    # legend
    for k, (label, color) in enumerate(perm_colors.items()):
        ax.add_patch(FancyBboxPatch((0.5+k*2.2, 0.2), 0.5, 0.4,
                     boxstyle="round,pad=0.05", facecolor=color,
                     edgecolor="white"))
        meaning = {"CRUD":"Full Access","CRU":"No Delete","R":"Read Only","—":"No Access"}
        ax.text(1.1+k*2.2, 0.4, f"{label} = {meaning[label]}",
                fontsize=7, va="center", color="#333333")

    return _save(fig, ax)


def flow_end_to_end():
    """End-to-End Cross-Module Flow"""
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 15); ax.set_ylim(0, 11)
    ax.set_facecolor(C["bg"])
    ax.set_title("Flow 9 – End-to-End Cross-Module Business Process",
                 fontsize=11, fontweight="bold", pad=8)

    # Module boxes (horizontal swim lanes concept, simplified as nodes)
    nodes = {
        "Settings\n(Locations/\nPlantTypes)": (1.0, 7.5),
        "Inventory\n(Seeds/\nChemicals)":     (1.0, 4.5),
        "Plantings":                          (4.5, 7.5),
        "Scouting":                           (4.5, 4.5),
        "Treatments":                         (4.5, 1.5),
        "Harvests":                           (8.0, 7.5),
        "Dispatch\nSlips":                    (8.0, 4.5),
        "Reservations":                       (8.0, 1.5),
        "Reports &\nDashboard":               (11.5, 4.5),
        "BOM /\nProduction":                  (11.5, 1.5),
    }
    node_colors = {
        "Settings\n(Locations/\nPlantTypes)": C["system"],
        "Inventory\n(Seeds/\nChemicals)":     C["data"],
        "Plantings":                          C["process"],
        "Scouting":                           C["decision"],
        "Treatments":                         C["decision"],
        "Harvests":                           C["process"],
        "Dispatch\nSlips":                    C["process"],
        "Reservations":                       C["system"],
        "Reports &\nDashboard":               C["start"],
        "BOM /\nProduction":                  C["data"],
    }

    centers = {}
    for label, (x, y) in nodes.items():
        cx, cy = _box(ax, x, y, 2.5, 1.5, node_colors[label], label,
                      shape="rounded", fontsize=7.5)
        centers[label] = (cx, cy)

    def ae(a, b, lbl="", **kw):
        x1, y1 = centers[a]
        x2, y2 = centers[b]
        _arrow(ax, x1, y1, x2, y2, label=lbl, **kw)

    ae("Settings\n(Locations/\nPlantTypes)", "Plantings", "defines")
    ae("Inventory\n(Seeds/\nChemicals)", "Plantings", "seed deduct")
    ae("Plantings", "Scouting", "batch ref")
    ae("Plantings", "Harvests", "stock calc")
    ae("Scouting", "Treatments", "triggers")
    ae("Treatments", "Inventory\n(Seeds/\nChemicals)", "chem deduct")
    ae("Harvests", "Dispatch\nSlips", "auto-fulfil")
    ae("Harvests", "Reservations", "conflict chk")
    ae("Reservations", "Plantings", "reserve qty")
    ae("Dispatch\nSlips", "Reservations", "customer")
    ae("Harvests", "Reports &\nDashboard", "data feed")
    ae("Reservations", "Reports &\nDashboard", "data feed")
    ae("Plantings", "BOM /\nProduction", "cost link")
    ae("Inventory\n(Seeds/\nChemicals)", "BOM /\nProduction", "price feed")
    ae("BOM /\nProduction", "Reports &\nDashboard", "profit data")

    return _save(fig, ax)


# ─── Word document builder ────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    return h


def add_flow(doc, title, buf, description):
    set_heading(doc, title, level=2)
    doc.add_paragraph(description)
    doc.add_picture(buf, width=Inches(6.2))
    doc.add_paragraph("")


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Cover ─────────────────────────────────────────────────
    title = doc.add_heading("Nursery Management App", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Business Process Flow Diagrams")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].bold = True
    doc.add_paragraph(
        "This document describes how each module interacts with others "
        "through end-to-end business processes. Each flow diagram uses the "
        "following shape conventions:"
    )

    # Legend table
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for cell, txt, shade in zip(
        hdr,
        ["Green oval = Start/End", "Blue rectangle = Process step",
         "Orange diamond = Decision", "Purple parallelogram = Data I/O",
         "Teal = System action"],
        ["D5F5E3", "D6EAF8", "FDEBD0", "E8DAEF", "D1F2EB"],
    ):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shade)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    doc.add_paragraph("")
    doc.add_page_break()

    # ── Section 1: System Setup ───────────────────────────────
    set_heading(doc, "1. System Setup & Onboarding", level=1)
    doc.add_paragraph(
        "Before any nursery operations can begin, the system must be configured "
        "with reference data — locations, plant types, inventory items, and user "
        "accounts. This flow shows the one-time setup sequence performed by an Admin."
    )
    add_flow(doc, "Flow 1 – System Setup & Onboarding", flow_onboarding(),
             "Admin configures all reference data before operational modules are used. "
             "If any data is missing, the loop returns to fill gaps.")
    doc.add_page_break()

    # ── Section 2: Planting Lifecycle ────────────────────────
    set_heading(doc, "2. Planting Lifecycle", level=1)
    doc.add_paragraph(
        "A planting batch is the core operational unit. It links plant types, "
        "locations, inventory (seeds), reservations, scouting reports, treatments, "
        "and harvests. This flow covers creation through to harvest-readiness."
    )
    add_flow(doc, "Flow 2 – Planting Lifecycle", flow_planting(),
             "Planting creation optionally deducts seeds from inventory, auto-generates "
             "a batch number, and checks for existing reservations. The batch is then "
             "monitored via scouting and treatments until harvest.")
    doc.add_page_break()

    # ── Section 3: Harvest & Dispatch ────────────────────────
    set_heading(doc, "3. Harvest & Dispatch Slip", level=1)
    doc.add_paragraph(
        "Two parallel paths exist for moving stock out of the nursery: "
        "a direct Harvest Record (used internally) and a Dispatch Slip "
        "(used by the dispatch team). The dispatch slip auto-creates a harvest "
        "if stock is available, or saves as pending for later fulfilment."
    )
    add_flow(doc, "Flow 3 – Harvest & Dispatch Slip", flow_harvest_dispatch(),
             "Left lane: standard harvest creation with reservation conflict handling. "
             "Right lane: dispatch slip flow — auto-fulfilled if stock exists, "
             "saved as pending otherwise. Pending slips can be edited and fulfilled later.")
    doc.add_page_break()

    # ── Section 4: Reservations ───────────────────────────────
    set_heading(doc, "4. Reservation Management", level=1)
    doc.add_paragraph(
        "Reservations allow customers to reserve seedlings from one or more "
        "planting batches before harvest. The system tracks availability in "
        "real time and prevents over-reservation. Completion records actual "
        "delivery quantities."
    )
    add_flow(doc, "Flow 4 – Reservation Management", flow_reservation(),
             "Customers can reserve across multiple batches. Each batch gets its own "
             "reservation record. On completion, actual delivered quantity is recorded "
             "and deducted from the planting's remaining quantity.")
    doc.add_page_break()

    # ── Section 5: Inventory ──────────────────────────────────
    set_heading(doc, "5. Inventory Management", level=1)
    doc.add_paragraph(
        "Inventory tracks chemicals, fertilisers, seeds, and supplies. "
        "Stock can be updated manually (purchase, usage, adjustment, waste) "
        "or automatically when a planting batch is created using seeds."
    )
    add_flow(doc, "Flow 5 – Inventory Management", flow_inventory(),
             "Left lane: manual stock transactions with low-stock alerts. "
             "Right lane: automatic seed deduction triggered by planting creation.")
    doc.add_page_break()

    # ── Section 6: Scouting & Treatment ──────────────────────
    set_heading(doc, "6. Scouting & Treatment", level=1)
    doc.add_paragraph(
        "Scouts regularly inspect the nursery and log observations (pests, "
        "diseases, nutrient deficiencies) against specific batches. Identified "
        "issues trigger treatment records that can be applied to individual "
        "batches or bulk-applied across a location."
    )
    add_flow(doc, "Flow 6 – Scouting & Treatment", flow_scouting_treatment(),
             "Scouting report captures observations with severity. If treatment is "
             "needed, a treatment record is created and linked to one or many plantings "
             "via the planting_treatments junction table. Chemical usage is deducted "
             "from inventory.")
    doc.add_page_break()

    # ── Section 7: BOM & Production ──────────────────────────
    set_heading(doc, "7. BOM & Production Cost Tracking", level=1)
    doc.add_paragraph(
        "The Bill of Materials (BOM) module allows managers to define the "
        "inputs required to produce a batch of seedlings and calculate the "
        "cost per unit. The Chemical Calculator supports mixing ratio decisions "
        "before treatments are applied."
    )
    add_flow(doc, "Flow 7 – BOM & Production Cost", flow_bom_production(),
             "BOM templates are created once, then assigned to planting batches. "
             "Costs are calculated by pulling current inventory prices. Results feed "
             "into the profit analysis reports.")
    doc.add_page_break()

    # ── Section 8: Permissions ────────────────────────────────
    set_heading(doc, "8. Role-Based Access Control", level=1)
    doc.add_paragraph(
        "Every module action is gated by the user's role. Four roles are "
        "defined: Admin, Manager, Staff, and Viewer. Permissions can be "
        "configured at the module level per role, or overridden per individual user."
    )
    add_flow(doc, "Flow 8 – Role-Based Access Control", flow_permissions(),
             "CRUD = full access. CRU = create/read/update (no delete). "
             "R = read only. — = no access. Admins have full access to all modules "
             "including user management.")
    doc.add_page_break()

    # ── Section 9: End-to-End ────────────────────────────────
    set_heading(doc, "9. End-to-End Cross-Module Process", level=1)
    doc.add_paragraph(
        "This diagram shows how all modules connect to each other in terms of "
        "data flow and process triggers. It is the high-level view of the entire "
        "system working together."
    )
    add_flow(doc, "Flow 9 – End-to-End Cross-Module Business Process",
             flow_end_to_end(),
             "Settings and Inventory are foundational. Plantings are the central "
             "operational record. Scouting → Treatments loop back to Inventory. "
             "Harvests, Reservations, and Dispatch Slips are the output flows. "
             "All modules feed Reports & Dashboard.")

    # ── Section 10: Interaction Summary table ────────────────
    doc.add_page_break()
    set_heading(doc, "10. Module Interaction Summary", level=1)
    doc.add_paragraph(
        "The table below summarises which modules trigger actions in other modules."
    )

    interactions = [
        ("Plantings",      "Inventory",    "Auto-deducts seed on creation"),
        ("Plantings",      "Reservations", "Batch qty drives reservation availability"),
        ("Plantings",      "Harvests",     "remaining_quantity updated on each harvest"),
        ("Plantings",      "Scouting",     "Batch linked to scouting reports"),
        ("Plantings",      "BOM",          "BOM template assigned per batch"),
        ("Scouting",       "Treatments",   "Identified issues trigger treatment creation"),
        ("Treatments",     "Inventory",    "Chemical usage deducted from stock"),
        ("Harvests",       "Dispatch Slips","Dispatch slip auto-creates harvest if stock OK"),
        ("Harvests",       "Reservations", "Checks & warns on reservation conflicts"),
        ("Reservations",   "Plantings",    "Completion deducts from remaining_quantity"),
        ("Dispatch Slips", "Harvests",     "Fulfilment creates a linked harvest record"),
        ("Inventory",      "BOM",          "Current prices used in cost calculations"),
        ("All Modules",    "Reports",      "All data aggregated in reports & dashboard"),
        ("Admin",          "All Modules",  "Role permissions gate every CRUD action"),
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for cell, hdr_txt in zip(tbl.rows[0].cells, ["From Module", "To Module", "Interaction"]):
        cell.text = hdr_txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for src, dst, desc in interactions:
        row = tbl.add_row().cells
        row[0].text = src
        row[1].text = dst
        row[2].text = desc
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(8)

    out = "/home/user/sg-51715866-84a2-4173-a0c3-6f30f6eab96a-1777618500/Process_Flows.docx"
    doc.save(out)
    print(f"✓ Document saved: {out}")


if __name__ == "__main__":
    build_document()

import os
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import numpy as np

OUTPUT_PATH = "/home/user/sg-51715866-84a2-4173-a0c3-6f30f6eab96a-1777618500/Nursery_App_Blueprint.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
C_HEADER    = RGBColor(0x1E, 0x40, 0xAF)   # blue-800
C_SUBHEAD   = RGBColor(0x16, 0x53, 0x4E)   # teal-800
C_ACCENT    = RGBColor(0x0F, 0x76, 0x6E)   # teal-600
C_TABLE_HDR = RGBColor(0x1E, 0x40, 0xAF)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT     = RGBColor(0xEF, 0xF6, 0xFF)

# ── UML colours ───────────────────────────────────────────────────────────────
UML_ENTITY  = "#DBEAFE"   # blue-100
UML_BORDER  = "#1E40AF"   # blue-800
UML_JUNC    = "#FEF3C7"   # amber-100
UML_JBORDER = "#D97706"
UML_LOOKUP  = "#DCFCE7"   # green-100
UML_LBORDER = "#16A34A"
UML_TEXT    = "#1E293B"
UML_ARROW   = "#374151"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.replace("#",""))
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = C_HEADER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = C_SUBHEAD
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = C_ACCENT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "#1E40AF")
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EFF6FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def img_to_docx(fig, doc, width=6.5):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close(fig)

# ══════════════════════════════════════════════════════════════════════════════
#  UML DIAGRAM GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def draw_entity(ax, x, y, w, h, name, fields, color=UML_ENTITY, border=UML_BORDER):
    """Draw a UML entity box with header + field list."""
    # Shadow
    shadow = FancyBboxPatch((x+0.04, y-0.04), w, h,
                             boxstyle="round,pad=0.02",
                             facecolor="#CBD5E1", edgecolor="none", zorder=1)
    ax.add_patch(shadow)
    # Body
    box = FancyBboxPatch((x, y), w, h,
                          boxstyle="round,pad=0.02",
                          facecolor=color, edgecolor=border,
                          linewidth=1.5, zorder=2)
    ax.add_patch(box)
    # Header band
    header = FancyBboxPatch((x, y + h - 0.28), w, 0.28,
                              boxstyle="round,pad=0.02",
                              facecolor=border, edgecolor=border,
                              linewidth=0, zorder=3)
    ax.add_patch(header)
    # Table name
    ax.text(x + w/2, y + h - 0.14, name,
            ha="center", va="center", fontsize=8, fontweight="bold",
            color="white", zorder=4)
    # Fields
    if fields:
        field_h = (h - 0.32) / max(len(fields), 1)
        for i, f in enumerate(fields):
            fy = y + h - 0.32 - (i + 0.6) * field_h
            style = "italic" if f.startswith("FK:") or f.startswith("PK:") else "normal"
            weight = "bold" if f.startswith("PK:") else "normal"
            color_f = "#1E40AF" if f.startswith("PK:") else \
                      "#6B21A8" if f.startswith("FK:") else UML_TEXT
            ax.text(x + 0.08, fy, f, ha="left", va="center",
                    fontsize=6.5, fontstyle=style, fontweight=weight,
                    color=color_f, zorder=4)

def draw_arrow(ax, x1, y1, x2, y2, label="", style="1:N"):
    """Draw a relationship arrow with crow's foot notation."""
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=UML_ARROW,
                                lw=1.2, mutation_scale=12),
                zorder=5)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.05, my+0.05, label, fontsize=6,
                color="#4B5563", zorder=6,
                bbox=dict(boxstyle="round,pad=0.1", facecolor="white",
                          edgecolor="none", alpha=0.8))

def make_fig(w, h):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, w)
    ax.set_ylim(0, h)
    ax.axis("off")
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")
    return fig, ax

def legend_patch(color, border, label):
    return mpatches.Patch(facecolor=color, edgecolor=border,
                          linewidth=1.2, label=label)

# ── Diagram 1: Core Growing Chain ─────────────────────────────────────────────
def uml_core_chain():
    fig, ax = make_fig(13, 7)
    ax.set_title("Core Growing Chain", fontsize=13, fontweight="bold",
                 color=UML_BORDER, pad=8)

    # plant_types
    draw_entity(ax, 0.2, 3.8, 2.4, 2.8, "plant_types",
                ["PK: id", "name", "variety", "growth_duration",
                 "germination_rate", "default_selling_price"])

    # locations
    draw_entity(ax, 0.2, 0.4, 2.4, 2.8, "locations",
                ["PK: id", "name", "type", "capacity", "current_occupancy"])

    # plantings  (centre)
    draw_entity(ax, 4.5, 2.2, 3.0, 3.4, "plantings",
                ["PK: id", "batch_number", "date_planted",
                 "expected_harvest_date", "quantity",
                 "remaining_quantity", "status",
                 "FK: plant_type_id", "FK: location_id",
                 "selling_price"], color="#BFDBFE")

    # harvests
    draw_entity(ax, 9.5, 4.6, 2.8, 2.0, "harvests",
                ["PK: id", "FK: planting_id", "harvest_date",
                 "quantity_harvested", "quality", "status"])

    # reservations
    draw_entity(ax, 9.5, 2.2, 2.8, 2.2, "reservations",
                ["PK: id", "FK: planting_id", "customer_name",
                 "quantity_reserved", "payment_status",
                 "collection_date", "final_quantity"])

    # dispatch_slips
    draw_entity(ax, 9.5, 0.2, 2.8, 1.8, "dispatch_slips",
                ["PK: id", "FK: planting_id", "FK: harvest_id",
                 "quantity_requested", "status", "customer_name"],
                color="#FEF3C7", border=UML_JBORDER)

    # Arrows
    draw_arrow(ax, 2.6, 5.2, 4.5, 3.8, "plant_type_id\n1:N")
    draw_arrow(ax, 2.6, 1.8, 4.5, 2.8, "location_id\n1:N")
    draw_arrow(ax, 7.5, 4.6, 9.5, 5.2, "planting_id\n1:N")
    draw_arrow(ax, 7.5, 3.9, 9.5, 3.3, "planting_id\n1:N")
    draw_arrow(ax, 7.5, 3.2, 9.5, 1.4, "planting_id\n1:N")
    draw_arrow(ax, 11.3, 1.8, 11.3, 4.6, "harvest_id\n1:1", style="1:1")

    legend = [legend_patch(UML_ENTITY, UML_BORDER, "Core table"),
              legend_patch("#BFDBFE", UML_BORDER, "Central hub (plantings)"),
              legend_patch(UML_JUNC, UML_JBORDER, "New feature (dispatch_slips)")]
    ax.legend(handles=legend, loc="lower left", fontsize=7,
              framealpha=0.9, edgecolor="#CBD5E1")
    fig.tight_layout()
    return fig

# ── Diagram 2: Scouting System ─────────────────────────────────────────────────
def uml_scouting():
    fig, ax = make_fig(13, 6)
    ax.set_title("Scouting System", fontsize=13, fontweight="bold",
                 color=UML_BORDER, pad=8)

    draw_entity(ax, 0.3, 1.5, 3.2, 3.6, "scouting_reports",
                ["PK: id", "FK: planting_id", "FK: created_by",
                 "scouting_date", "scout_name", "crop_type",
                 "greenhouse_location", "overall_health_rating",
                 "weather_conditions"], color="#BFDBFE")

    draw_entity(ax, 5.5, 3.8, 2.8, 1.8, "scouting_pests",
                ["PK: id", "FK: report_id", "pest_name",
                 "severity", "percent_trays_affected"])

    draw_entity(ax, 5.5, 1.8, 2.8, 1.8, "scouting_diseases",
                ["PK: id", "FK: report_id", "disease_name",
                 "severity", "recommended_action"])

    draw_entity(ax, 5.5, 0.0, 2.8, 1.6, "scouting_nutrients",
                ["PK: id", "FK: report_id", "symptom",
                 "suspected_deficiency", "severity"])

    draw_entity(ax, 9.8, 3.8, 2.8, 1.4, "scouting_pest_types",
                ["PK: id", "name", "is_active", "display_order"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 9.8, 2.0, 2.8, 1.4, "scouting_disease_types",
                ["PK: id", "name", "is_active", "display_order"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 9.8, 0.2, 2.8, 1.4, "scouting_nutrient_types",
                ["PK: id", "name", "is_active", "display_order"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_arrow(ax, 3.5, 3.3, 5.5, 4.5, "report_id\n1:N")
    draw_arrow(ax, 3.5, 3.0, 5.5, 2.8, "report_id\n1:N")
    draw_arrow(ax, 3.5, 2.7, 5.5, 0.8, "report_id\n1:N")
    draw_arrow(ax, 8.3, 4.6, 9.8, 4.5, "lookup")
    draw_arrow(ax, 8.3, 2.8, 9.8, 2.7, "lookup")
    draw_arrow(ax, 8.3, 0.8, 9.8, 0.9, "lookup")

    legend = [legend_patch(UML_ENTITY, UML_BORDER, "Report table"),
              legend_patch("#BFDBFE", UML_BORDER, "Parent (scouting_reports)"),
              legend_patch(UML_LOOKUP, UML_LBORDER, "Configurable lookup")]
    ax.legend(handles=legend, loc="lower left", fontsize=7,
              framealpha=0.9, edgecolor="#CBD5E1")
    fig.tight_layout()
    return fig

# ── Diagram 3: Production / BOM Chain ─────────────────────────────────────────
def uml_production():
    fig, ax = make_fig(13, 6.5)
    ax.set_title("Production / BOM Chain", fontsize=13, fontweight="bold",
                 color=UML_BORDER, pad=8)

    draw_entity(ax, 0.2, 3.8, 2.6, 1.8, "plant_types",
                ["PK: id", "name", "variety"])

    draw_entity(ax, 0.2, 1.2, 2.6, 2.2, "profiles",
                ["PK: id", "full_name", "email", "role"])

    draw_entity(ax, 4.2, 2.5, 3.0, 3.4, "bom_templates",
                ["PK: id", "name", "FK: plant_type_id",
                 "FK: created_by", "base_batch_size",
                 "estimated_success_rate",
                 "target_selling_price", "status"], color="#BFDBFE")

    draw_entity(ax, 4.2, 0.2, 3.0, 2.0, "planting_bom_costs",
                ["PK: id", "FK: planting_id (1:1)",
                 "FK: bom_template_id",
                 "estimated_total_cost",
                 "actual_total_cost", "cost_variance"])

    draw_entity(ax, 9.0, 3.4, 2.8, 2.8, "bom_items",
                ["PK: id", "FK: template_id",
                 "FK: category_id", "FK: inventory_item_id",
                 "quantity_value", "quantity_formula",
                 "item_type"])

    draw_entity(ax, 9.0, 1.6, 2.8, 1.6, "bom_categories",
                ["PK: id", "name", "color", "sort_order"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 9.0, 0.0, 2.8, 1.4, "formula_templates",
                ["PK: id", "name", "formula", "variables"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_arrow(ax, 2.8, 4.6, 4.2, 4.2, "plant_type_id\n1:N")
    draw_arrow(ax, 2.8, 2.2, 4.2, 3.2, "created_by\n1:N")
    draw_arrow(ax, 7.2, 3.8, 9.0, 4.8, "template_id\n1:N")
    draw_arrow(ax, 5.7, 2.5, 5.7, 2.2, "bom_template_id\n1:N")
    draw_arrow(ax, 11.4, 3.4, 11.4, 3.2, "category_id\n1:N")

    legend = [legend_patch(UML_ENTITY, UML_BORDER, "Entity"),
              legend_patch("#BFDBFE", UML_BORDER, "Central (bom_templates)"),
              legend_patch(UML_LOOKUP, UML_LBORDER, "Lookup / standalone")]
    ax.legend(handles=legend, loc="lower right", fontsize=7,
              framealpha=0.9, edgecolor="#CBD5E1")
    fig.tight_layout()
    return fig

# ── Diagram 4: Inventory & Chemicals ──────────────────────────────────────────
def uml_inventory():
    fig, ax = make_fig(13, 6)
    ax.set_title("Inventory & Chemical Chain", fontsize=13, fontweight="bold",
                 color=UML_BORDER, pad=8)

    draw_entity(ax, 3.5, 2.0, 3.0, 3.4, "inventory_items",
                ["PK: id", "name", "current_stock",
                 "minimum_stock", "unit_price",
                 "category", "unit_of_measure"], color="#BFDBFE")

    draw_entity(ax, 0.2, 3.2, 2.6, 1.6, "inventory_categories",
                ["PK: id", "name", "color"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 0.2, 1.2, 2.6, 1.6, "inventory_units",
                ["PK: id", "name", "abbreviation", "type"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 0.2, 0.0, 2.6, 1.0, "inventory_suppliers",
                ["PK: id", "name", "contact_person", "email"],
                color=UML_LOOKUP, border=UML_LBORDER)

    draw_entity(ax, 8.0, 3.2, 2.8, 2.4, "stock_transactions",
                ["PK: id", "FK: item_id", "FK: created_by",
                 "transaction_type", "quantity",
                 "unit_price", "reference_type"])

    draw_entity(ax, 8.0, 1.0, 2.8, 2.0, "chemical_products",
                ["PK: id", "FK: inventory_item_id",
                 "name", "type", "ec_factor",
                 "npk_n/p/k", "recommended_concentration"])

    draw_entity(ax, 8.0, 0.0, 2.8, 0.9, "saved_mixes",
                ["PK: id", "FK: product_id",
                 "concentration", "calculated_ec"],
                color=UML_JUNC, border=UML_JBORDER)

    draw_arrow(ax, 3.5, 3.8, 2.8, 4.0, "category\n(text ref)")
    draw_arrow(ax, 3.5, 3.4, 2.8, 2.0, "unit_of_measure\n(text ref)")
    draw_arrow(ax, 6.5, 4.2, 8.0, 4.2, "item_id\n1:N")
    draw_arrow(ax, 6.5, 3.0, 8.0, 2.0, "inventory_item_id\n1:N")
    draw_arrow(ax, 9.4, 1.0, 9.4, 0.9, "product_id\n1:N")

    legend = [legend_patch(UML_ENTITY, UML_BORDER, "Entity"),
              legend_patch("#BFDBFE", UML_BORDER, "Central (inventory_items)"),
              legend_patch(UML_LOOKUP, UML_LBORDER, "Lookup / config"),
              legend_patch(UML_JUNC, UML_JBORDER, "Saved data")]
    ax.legend(handles=legend, loc="lower left", fontsize=7,
              framealpha=0.9, edgecolor="#CBD5E1")
    fig.tight_layout()
    return fig

# ── Diagram 5: Access Control ──────────────────────────────────────────────────
def uml_access_control():
    fig, ax = make_fig(13, 5.5)
    ax.set_title("Access Control & Permissions", fontsize=13, fontweight="bold",
                 color=UML_BORDER, pad=8)

    draw_entity(ax, 0.2, 1.5, 2.8, 2.8, "profiles",
                ["PK: id", "full_name", "email", "username",
                 "role: admin|manager|staff|viewer",
                 "avatar_url"], color="#BFDBFE")

    draw_entity(ax, 4.5, 2.2, 2.8, 2.2, "permissions",
                ["PK: id", "module", "action", "resource",
                 "description"])

    draw_entity(ax, 4.5, 0.1, 2.8, 1.8, "password_history",
                ["PK: id", "FK: user_id",
                 "password_hash", "created_at"],
                color=UML_JUNC, border=UML_JBORDER)

    draw_entity(ax, 8.8, 2.8, 3.0, 1.8, "role_permissions",
                ["PK: id", "FK: permission_id",
                 "role", "can_create", "can_read",
                 "can_update", "can_delete"])

    draw_entity(ax, 8.8, 0.6, 3.0, 2.0, "user_permissions",
                ["PK: id", "FK: permission_id",
                 "FK: user_id", "can_create", "can_read",
                 "can_update", "can_delete"])

    draw_arrow(ax, 3.0, 3.0, 4.5, 3.5, "1:N per module")
    draw_arrow(ax, 3.0, 2.2, 4.5, 1.2, "user_id\n1:N")
    draw_arrow(ax, 7.3, 3.4, 8.8, 3.6, "permission_id\n1:N")
    draw_arrow(ax, 7.3, 3.0, 8.8, 1.8, "permission_id\n1:N")
    draw_arrow(ax, 3.0, 2.5, 8.8, 1.5, "user_id\n1:N")

    legend = [legend_patch(UML_ENTITY, UML_BORDER, "Entity"),
              legend_patch("#BFDBFE", UML_BORDER, "profiles (central)"),
              legend_patch(UML_JUNC, UML_JBORDER, "Audit / history")]
    ax.legend(handles=legend, loc="lower right", fontsize=7,
              framealpha=0.9, edgecolor="#CBD5E1")
    fig.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Nursery Management System")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = C_HEADER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Application Blueprint & Database Schematic")
    r2.font.size = Pt(16); r2.font.color.rgb = C_SUBHEAD

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Version 1.0  ·  June 2026")
    r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── 1. Overview ───────────────────────────────────────────────────────────
    add_heading(doc, "1. Application Overview")
    add_body(doc,
        "The Nursery Management System is a full-stack web application built on "
        "Next.js 15 (Pages Router) with TypeScript and Supabase (PostgreSQL). "
        "It covers the complete operational lifecycle of a plant nursery: from "
        "seeding and growing through pest monitoring, inventory, customer "
        "reservations, and dispatch.")
    add_body(doc,
        "The application is role-based (admin, manager, staff, viewer) and "
        "enforces permissions both in the UI layer (usePermissions hook) and at "
        "the database level (role_permissions table).")

    doc.add_paragraph()
    add_heading(doc, "Technology Stack", level=2)
    add_styled_table(doc,
        ["Layer", "Technology"],
        [["Frontend Framework", "Next.js 15 (Pages Router) + TypeScript"],
         ["UI Components",     "shadcn/ui (Radix UI primitives) + Tailwind CSS"],
         ["Backend / Database","Supabase (PostgreSQL 15)"],
         ["Authentication",    "Supabase Auth with custom profiles table"],
         ["Charts",            "Recharts + Chart.js"],
         ["Forms",             "React Hook Form + Zod"],
         ["Drag & Drop",       "dnd-kit"],
         ["PDF / Print",       "jsPDF + html2canvas + browser print API"],
         ["Payments",          "Stripe"]],
        col_widths=[2.0, 4.5])

    doc.add_page_break()

    # ── 2. Pages ──────────────────────────────────────────────────────────────
    add_heading(doc, "2. Pages & Modules")
    add_styled_table(doc,
        ["Route", "Module", "Description"],
        [["/", "Auth", "Login page with first-time admin initialisation"],
         ["/dashboard", "Dashboard", "Live KPI metrics: active batches, harvests, reservations, locations"],
         ["/plantings", "Growing", "Manage all plant batches — location, quantity, status, dates"],
         ["/locations", "Growing", "Greenhouse zones with capacity and occupancy tracking"],
         ["/plant-types", "Growing", "Species catalogue with pricing, growth duration, germination rates"],
         ["/harvests", "Harvest", "Log harvests + Dispatch Slips tab (pending/fulfilled/cancelled)"],
         ["/harvests/bulk", "Harvest", "Bulk-harvest multiple batches simultaneously"],
         ["/treatments", "Care", "Record chemical/biological treatments per batch"],
         ["/scouting", "Care", "Pest, disease & nutrient monitoring reports"],
         ["/inventory", "Stock", "Supplies, current stock, low-stock alerts, transaction history"],
         ["/reservations", "Sales", "Customer pre-orders with payment and collection tracking"],
         ["/production/bom", "Production", "Bill of Materials templates with cost breakdowns"],
         ["/production/calculator", "Production", "Chemical/nutrient EC mix calculator"],
         ["/reports/*", "Reports", "7 reports: yields, upcoming harvests, customer availability, location utilisation, movement analysis"],
         ["/admin/user-management", "Admin", "Create/edit users and assign roles"],
         ["/admin/roles-permissions", "Admin", "Configure per-role CRUD permission matrix"],
         ["/settings/*", "Settings", "Inventory categories, scouting types, production config"]],
        col_widths=[1.8, 1.4, 3.4])

    doc.add_page_break()

    # ── 3. Services ───────────────────────────────────────────────────────────
    add_heading(doc, "3. Service Layer")
    add_body(doc,
        "All database access goes through typed service files in src/services/. "
        "No raw queries appear in page components. Each service maps to one or "
        "more Supabase tables and exposes typed async functions.")
    add_styled_table(doc,
        ["Service File", "Responsibility"],
        [["authService.ts",               "Login, logout, profile management, password reset"],
         ["plantingService.ts",           "CRUD for plantings, batch tracking, remaining quantity"],
         ["harvestService.ts",            "Harvest recording, validation against reservations, bulk harvests"],
         ["locationService.ts",           "Location CRUD, capacity and occupancy tracking"],
         ["plantTypeService.ts",          "Species/variety catalogue management"],
         ["treatmentService.ts",          "Treatment logging, planting associations"],
         ["scoutingService.ts",           "Scouting reports + pest/disease/nutrient findings"],
         ["reservationService.ts",        "Customer reservations, payment status, quantity management"],
         ["inventoryService.ts",          "Stock CRUD, low-stock alerts, transaction history"],
         ["inventorySettingsService.ts",  "Inventory category and unit configuration"],
         ["scoutingSettingsService.ts",   "Pest/disease/nutrient type configuration"],
         ["rolePermissionService.ts",     "Role-permission matrix, CRUD access control"],
         ["adminService.ts",              "Admin user operations, system initialisation"],
         ["bomService.ts",                "BOM template management, cost calculations"],
         ["chemicalCalculatorService.ts", "EC/nutrient solution calculations"],
         ["dispatchSlipService.ts",       "Dispatch slip creation, auto-fulfilment, print support"]],
        col_widths=[2.4, 4.2])

    doc.add_page_break()

    # ── 4. Data Model ─────────────────────────────────────────────────────────
    add_heading(doc, "4. Database Tables — Complete Field Reference")

    add_heading(doc, "4.1 Core Growing Tables", level=2)
    add_styled_table(doc,
        ["Table", "Key Fields", "References"],
        [["plant_types",   "id, name, variety, growth_duration, germination_rate, default_selling_price", "—"],
         ["locations",     "id, name, type, capacity, current_occupancy", "—"],
         ["plantings",     "id, batch_number, date_planted, expected_harvest_date, quantity, remaining_quantity, status, selling_price", "plant_types, locations"],
         ["harvests",      "id, harvest_date, quantity_harvested, quality, status, is_closed, notes", "plantings"],
         ["reservations",  "id, customer_name, customer_phone, customer_email, quantity_reserved, payment_status, collection_date, final_quantity, total_amount", "plantings"],
         ["dispatch_slips","id, quantity_requested, dispatch_date, customer_name, destination, status (pending|fulfilled|cancelled)", "plantings, harvests"],
         ["treatments",    "id, name, type, dosage, application_method, application_date, applied_by, planting_ids[]", "—"],
         ["planting_treatments", "planting_id, treatment_id (junction)", "plantings, treatments"],
         ["planting_bom_costs", "id, estimated_*/actual_* costs per category, cost_variance", "plantings (1:1)"]],
        col_widths=[1.8, 3.2, 1.6])

    doc.add_paragraph()
    add_heading(doc, "4.2 Scouting Tables", level=2)
    add_styled_table(doc,
        ["Table", "Key Fields", "References"],
        [["scouting_reports",      "id, scouting_date, scout_name, crop_type, greenhouse_location, overall_health_rating, weather_conditions", "plantings, profiles"],
         ["scouting_pests",        "id, pest_name, severity, percent_trays_affected, action_required, distribution_pattern", "scouting_reports"],
         ["scouting_diseases",     "id, disease_name, severity, percent_trays_affected, recommended_action", "scouting_reports"],
         ["scouting_nutrients",    "id, symptom, suspected_deficiency, severity, percent_affected", "scouting_reports"],
         ["scouting_pest_types",   "id, name, is_active, display_order", "—"],
         ["scouting_disease_types","id, name, is_active, display_order", "—"],
         ["scouting_nutrient_types","id, name, is_active, display_order", "—"],
         ["scouting_actions",      "id, name, category, is_active, display_order", "—"]],
        col_widths=[2.0, 3.4, 1.2])

    doc.add_paragraph()
    add_heading(doc, "4.3 Production & BOM Tables", level=2)
    add_styled_table(doc,
        ["Table", "Key Fields", "References"],
        [["bom_templates",     "id, name, base_batch_size, estimated_success_rate, target_selling_price, status", "plant_types, profiles"],
         ["bom_items",         "id, item_type, quantity_value, quantity_formula, custom_name, custom_unit", "bom_templates, bom_categories, inventory_items"],
         ["bom_categories",    "id, name, color, sort_order", "—"],
         ["bom_headers",       "id, name, product_name, version, status", "—"],
         ["formula_templates", "id, name, formula, variables[]", "—"],
         ["chemical_products", "id, name, type, form, ec_factor, npk_n/p/k, min/max/recommended_concentration", "inventory_items"],
         ["saved_mixes",       "id, concentration, water_volume, calculated_ec, target_ec, applied_to_planting_ids[]", "chemical_products, profiles"]],
        col_widths=[2.0, 3.4, 1.6])

    doc.add_paragraph()
    add_heading(doc, "4.4 Inventory Tables", level=2)
    add_styled_table(doc,
        ["Table", "Key Fields", "References"],
        [["inventory_items",      "id, name, current_stock, minimum_stock, unit_price, category, unit_of_measure", "—"],
         ["inventory_categories", "id, name, color, description", "—"],
         ["inventory_units",      "id, name, abbreviation, type", "—"],
         ["inventory_suppliers",  "id, name, contact_person, email, phone, address", "—"],
         ["stock_transactions",   "id, transaction_type, quantity, unit_price, total_cost, reference_type, reference_id", "inventory_items, profiles"]],
        col_widths=[2.0, 3.6, 1.0])

    doc.add_paragraph()
    add_heading(doc, "4.5 Access Control Tables", level=2)
    add_styled_table(doc,
        ["Table", "Key Fields", "References"],
        [["profiles",          "id, full_name, email, username, role (admin|manager|staff|viewer), avatar_url", "—"],
         ["permissions",       "id, module, action, resource, description", "—"],
         ["role_permissions",  "id, role, can_create, can_read, can_update, can_delete", "permissions"],
         ["user_permissions",  "id, user_id, can_create, can_read, can_update, can_delete", "profiles, permissions"],
         ["password_history",  "id, user_id, password_hash, created_at", "—"]],
        col_widths=[2.0, 3.4, 1.2])

    doc.add_page_break()

    # ── 5. UML Diagrams ───────────────────────────────────────────────────────
    add_heading(doc, "5. UML Entity-Relationship Diagrams")
    add_body(doc,
        "The following diagrams show all foreign-key relationships between "
        "tables. Blue boxes are primary entities, blue-shaded boxes are "
        "central hubs, green boxes are lookup/configuration tables, and "
        "amber boxes are junction or feature tables.")

    doc.add_paragraph()
    add_heading(doc, "5.1  Core Growing Chain", level=2)
    add_body(doc, "plant_types and locations feed into plantings, which is the "
             "central hub for harvests, reservations, and dispatch slips.")
    img_to_docx(uml_core_chain(), doc, width=6.4)

    doc.add_paragraph()
    add_heading(doc, "5.2  Scouting System", level=2)
    add_body(doc, "scouting_reports links to plantings and spawns three child "
             "finding tables (pests, diseases, nutrients) and four configurable "
             "lookup lists.")
    img_to_docx(uml_scouting(), doc, width=6.4)

    doc.add_page_break()
    add_heading(doc, "5.3  Production / BOM Chain", level=2)
    add_body(doc, "bom_templates connect plant types to costed item lists "
             "(bom_items) and track estimated vs actual spend per planting "
             "(planting_bom_costs).")
    img_to_docx(uml_production(), doc, width=6.4)

    doc.add_paragraph()
    add_heading(doc, "5.4  Inventory & Chemical Chain", level=2)
    add_body(doc, "inventory_items is the central stock table. stock_transactions "
             "provides an audit trail; chemical_products and saved_mixes support "
             "the nutrient calculator.")
    img_to_docx(uml_inventory(), doc, width=6.4)

    doc.add_page_break()
    add_heading(doc, "5.5  Access Control & Permissions", level=2)
    add_body(doc, "profiles holds every user and their role. The permissions "
             "table defines per-module actions; role_permissions sets defaults "
             "per role while user_permissions allows individual overrides.")
    img_to_docx(uml_access_control(), doc, width=6.4)

    doc.add_page_break()

    # ── 6. Relationship Summary ───────────────────────────────────────────────
    add_heading(doc, "6. Full Relationship Summary")
    add_styled_table(doc,
        ["Child Table", "Parent Table", "Foreign Key", "Cardinality"],
        [["plantings",          "plant_types",       "plant_type_id",       "N:1"],
         ["plantings",          "locations",         "location_id",         "N:1"],
         ["harvests",           "plantings",         "planting_id",         "N:1"],
         ["reservations",       "plantings",         "planting_id",         "N:1"],
         ["dispatch_slips",     "plantings",         "planting_id",         "N:1"],
         ["dispatch_slips",     "harvests",          "harvest_id",          "1:1"],
         ["planting_treatments","plantings",         "planting_id",         "M:N junction"],
         ["planting_treatments","treatments",        "treatment_id",        "M:N junction"],
         ["planting_bom_costs", "plantings",         "planting_id",         "1:1"],
         ["scouting_reports",   "plantings",         "planting_id",         "N:1 (optional)"],
         ["scouting_reports",   "profiles",          "created_by",          "N:1"],
         ["scouting_pests",     "scouting_reports",  "report_id",           "N:1"],
         ["scouting_diseases",  "scouting_reports",  "report_id",           "N:1"],
         ["scouting_nutrients", "scouting_reports",  "report_id",           "N:1"],
         ["bom_templates",      "plant_types",       "plant_type_id",       "N:1"],
         ["bom_templates",      "profiles",          "created_by",          "N:1"],
         ["bom_items",          "bom_templates",     "template_id",         "N:1"],
         ["bom_items",          "bom_categories",    "category_id",         "N:1"],
         ["bom_items",          "inventory_items",   "inventory_item_id",   "N:1"],
         ["chemical_products",  "inventory_items",   "inventory_item_id",   "N:1"],
         ["saved_mixes",        "chemical_products", "product_id",          "N:1"],
         ["saved_mixes",        "profiles",          "created_by",          "N:1"],
         ["stock_transactions", "inventory_items",   "item_id",             "N:1"],
         ["stock_transactions", "profiles",          "created_by",          "N:1"],
         ["role_permissions",   "permissions",       "permission_id",       "N:1"],
         ["user_permissions",   "permissions",       "permission_id",       "N:1"],
         ["user_permissions",   "profiles",          "user_id",             "N:1"]],
        col_widths=[1.8, 1.8, 1.8, 1.2])

    doc.add_page_break()

    # ── 7. Permissions Matrix ─────────────────────────────────────────────────
    add_heading(doc, "7. Role Permissions Matrix")
    add_styled_table(doc,
        ["Module", "Admin", "Manager", "Staff", "Viewer"],
        [["Plantings",    "CRUD", "CRU",  "CRU",  "R"],
         ["Locations",    "CRUD", "CRU",  "R",    "R"],
         ["Plant Types",  "CRUD", "CRU",  "R",    "R"],
         ["Harvests",     "CRUD", "CRU",  "CRU",  "R"],
         ["Dispatch Slips","CRUD","CRU",  "CRU",  "R"],
         ["Treatments",   "CRUD", "CRU",  "CRU",  "R"],
         ["Scouting",     "CRUD", "CRU",  "CRU",  "R"],
         ["Inventory",    "CRUD", "CRU",  "CRU",  "R"],
         ["Reservations", "CRUD", "CRU",  "CRU",  "R"],
         ["Production",   "CRUD", "CRU",  "R",    "R"],
         ["Reports",      "CRUD", "R",    "R",    "R"],
         ["Admin",        "CRUD", "—",    "—",    "—"],
         ["Settings",     "CRUD", "CRU",  "—",    "—"]],
        col_widths=[2.0, 1.0, 1.0, 1.0, 1.0])

    add_body(doc, "C = Create  ·  R = Read  ·  U = Update  ·  D = Delete")

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"✓ Document saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_document()
